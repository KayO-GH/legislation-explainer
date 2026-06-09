"""Utilities for ingesting and normalizing documents."""

from __future__ import annotations

import io
import os
import re
import tempfile
from urllib.parse import urlparse
from typing import Iterable

import requests
from docx import Document as DocxDocument
from lxml import html
from readability import Document as ReadabilityDocument
from pypdf import PdfReader

from config import MAX_UPLOAD_SIZE_MB, TIMEOUT_SECONDS

SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


class IngestionError(RuntimeError):
    pass


def _ensure_size_limit(file_size: int) -> None:
    max_bytes = MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if file_size > max_bytes:
        raise IngestionError(
            f"File exceeds maximum size of {MAX_UPLOAD_SIZE_MB} MB."
        )


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    text_chunks: list[str] = []
    for page in reader.pages:
        chunk = page.extract_text() or ""
        text_chunks.append(chunk.strip())
    return _normalize_text(_repair_known_collapsed_terms("\n".join(text_chunks)))


def extract_text_from_docx(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        document = DocxDocument(tmp_path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return _normalize_text("\n".join(paragraphs))
    finally:
        os.unlink(tmp_path)


def extract_text_from_text(data: bytes) -> str:
    return _normalize_text(data.decode("utf-8", errors="ignore"))


def ingest_file(filename: str, file_bytes: bytes) -> str:
    _ensure_size_limit(len(file_bytes))
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise IngestionError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    return extract_text_from_text(file_bytes)


def fetch_url_content(url: str, timeout_seconds: int = TIMEOUT_SECONDS) -> str:
    response = requests.get(url, timeout=timeout_seconds)
    response.raise_for_status()
    content_type = (response.headers.get("content-type") or "").lower()
    path = urlparse(url).path.lower()

    if "application/pdf" in content_type or path.endswith(".pdf"):
        return extract_text_from_pdf(response.content)
    if (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type
        or path.endswith(".docx")
    ):
        return extract_text_from_docx(response.content)
    if any(text_type in content_type for text_type in ("text/plain", "text/markdown")) or any(
        path.endswith(ext) for ext in (".txt", ".md")
    ):
        return extract_text_from_text(response.content)

    readable = ReadabilityDocument(_strip_xml_control_chars(response.text))
    summary_html = _strip_xml_control_chars(readable.summary(html_partial=True))
    parsed = html.fromstring(summary_html)
    title = readable.short_title() or readable.title() or ""
    text_content = parsed.text_content()
    return _normalize_text("\n".join([title, text_content]))


def combine_sources(sources: Iterable[str]) -> str:
    return _normalize_text("\n\n".join(filter(None, sources)))


def _normalize_text(value: str) -> str:
    value = _strip_xml_control_chars(value)
    clean_lines = [line.strip() for line in value.splitlines()]
    return "\n".join(line for line in clean_lines if line)


def _strip_xml_control_chars(value: str) -> str:
    return _CONTROL_CHARS_RE.sub("", value)


def _repair_known_collapsed_terms(value: str) -> str:
    replacements = {
        "NATIONALINFORMATIONTECHNOLOGYAUTHORITY": "National Information Technology Authority",
        "NationalInformationTechnologyAuthority": "National Information Technology Authority",
    }
    for collapsed, expanded in replacements.items():
        value = value.replace(collapsed, expanded)
    return value
